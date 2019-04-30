import os, sys
from docx import Document
from docxcompose.composer import Composer

def read_docx(fp):
    f = open(fp, 'rb')
    docx = Document(f)
    f.close()
    return docx

def head(fp):    
    document = Document()
    document.add_heading(os.path.basename(fp).replace('docx', 'py'), 0)
    return document

def concat_docx(dir_p):
    docxs = [os.path.join(dir_p, fn) for fn in os.listdir(dir_p) if fn.endswith('docx')]
    up_bound = len(docxs)
    d0 = read_docx(docxs[0])
    d0.add_page_break()
    h0 = head(docxs[0])
    compose = Composer(h0)
    compose.append(d0)
    for count, docx in enumerate(docxs[1:], start=2):
        sub_head = head(docx)
        compose.append(sub_head)
        sub_d = read_docx(docx)
        if count < up_bound:
            sub_d.add_page_break()
        compose.append(sub_d)
    compose.save('c.docx')

if __name__ == '__main__':
	concat_docx(sys.argv[1])