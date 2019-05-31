import os, sys
from docx import Document
from docxcompose.composer import Composer

def read_docx(fp):
    f = open(fp, 'rb')
    docx = Document(f)
    f.close()
    return docx

def head(fp, suffix):    
    document = Document()
    document.add_heading(os.path.basename(fp).replace('docx', suffix), 0)
    return document

def finder(dir_p, ends='docx'):
	return [os.path.join(dir_p, fn) for fn in os.listdir(dir_p) if fn.endswith(ends)]

def concat_docx(dir_p, dst_p, suffix, name='codes'):
    docxs = finder(dir_p)
    up_bound = len(docxs)
    d0 = read_docx(docxs[0])
    d0.add_page_break()
    h0 = head(docxs[0], suffix)
    compose = Composer(h0)
    compose.append(d0)
    for count, docx in enumerate(docxs[1:], start=2):
        sub_head = head(docx, suffix)
        compose.append(sub_head)
        sub_d = read_docx(docx)
        if count < up_bound:
            sub_d.add_page_break()
        compose.append(sub_d)
    compose.save(os.path.join(dst_p, f'{name}.docx'))

if __name__ == '__main__':
	concat_docx(sys.argv[1], sys.argv[2], sys.argv[3])
	# files = finder(sys.argv[1], ends='doc')
	# compose = Composer(read_docx(files[0]))
	# concat_core(files[1], compose, start=2)